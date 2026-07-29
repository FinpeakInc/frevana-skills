import csv
import ctypes
import hashlib
import html
import importlib
import importlib.util
import os
import re
import signal
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from core.config import (
    DEFAULT_MODE,
    IMAGE_EXTS,
    OPTIONAL_EXTS,
    SKIP_DIRS,
    SKIP_FILE_NAMES,
    SKIP_FILE_SUFFIXES,
    TEXT_EXTS,
)
from core.errors import LocalKnowledgeError


LEGACY_DOC_PARSERS = ("textutil", "antiword", "catdoc")
WORD_COM_HELPER = Path(__file__).with_name("word_com.py")
PYTHON_DOCUMENT_PARSERS = (
    (".pdf", "pypdf"),
    (".docx", "docx"),
    (".xlsx", "openpyxl"),
)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def is_binary_sample(path: Path) -> bool:
    try:
        with path.open("rb") as f:
            sample = f.read(4096)
    except OSError:
        return True
    return b"\x00" in sample


def safe_read_text(path: Path) -> str:
    data = path.read_bytes()
    return decode_text(data)


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def find_legacy_doc_parser() -> Optional[Tuple[str, str]]:
    parsers = find_legacy_doc_parsers()
    return parsers[0] if parsers else None


def find_legacy_doc_parsers() -> List[Tuple[str, str]]:
    parsers: List[Tuple[str, str]] = []
    if word_com_available():
        parsers.append(("word-com", sys.executable))
    for name in LEGACY_DOC_PARSERS:
        executable = shutil.which(name)
        if executable:
            parsers.append((name, executable))
    return parsers


def word_com_available() -> bool:
    if sys.platform != "win32":
        return False
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore  # noqa: F401

        pythoncom.CLSIDFromProgID("Word.Application")
    except Exception:
        return False
    return True


def legacy_doc_parser_name() -> Optional[str]:
    parser = find_legacy_doc_parser()
    return parser[0] if parser else None


def legacy_doc_parser_signature() -> List[Dict[str, Any]]:
    signature: List[Dict[str, Any]] = []
    for name, executable in find_legacy_doc_parsers():
        item: Dict[str, Any] = {"name": name, "path": executable}
        if name == "word-com":
            try:
                spec = importlib.util.find_spec("win32com.client")
            except (ImportError, ModuleNotFoundError, ValueError):
                spec = None
            origin = spec.origin if spec is not None else None
            if origin:
                item["module_path"] = origin
                try:
                    module_stat = Path(origin).stat()
                except OSError:
                    pass
                else:
                    item["module_size"] = module_stat.st_size
                    item["module_mtime_ns"] = module_stat.st_mtime_ns
        try:
            stat = Path(executable).stat()
        except OSError:
            pass
        else:
            item["size"] = stat.st_size
            item["mtime_ns"] = stat.st_mtime_ns
        signature.append(item)
    return signature


def python_document_parser_signature() -> List[Dict[str, Any]]:
    signature: List[Dict[str, Any]] = []
    for extension, module in PYTHON_DOCUMENT_PARSERS:
        item: Dict[str, Any] = {"extension": extension, "module": module}
        try:
            importlib.import_module(module)
        except Exception:
            item["import_ok"] = False
        else:
            item["import_ok"] = True
        try:
            spec = importlib.util.find_spec(module)
        except (ImportError, ValueError):
            spec = None
        origin = spec.origin if spec is not None else None
        if origin:
            item["path"] = origin
            try:
                stat = Path(origin).stat()
            except OSError:
                pass
            else:
                item["size"] = stat.st_size
                item["mtime_ns"] = stat.st_mtime_ns
        signature.append(item)
    return signature


def legacy_doc_command(
    name: str,
    executable: str,
    path: Path,
    pid_file: Optional[Path] = None,
) -> List[str]:
    if name == "word-com":
        if pid_file is None:
            raise LocalKnowledgeError("Word COM parser requires a PID file.")
        return [executable, str(WORD_COM_HELPER), str(path), str(pid_file)]
    if name == "textutil":
        return [executable, "-convert", "txt", "-stdout", str(path)]
    return [executable, str(path)]


def terminate_word_process(pid_file: Path) -> None:
    if sys.platform != "win32" or not pid_file.is_file():
        return
    try:
        pid = int(pid_file.read_text(encoding="ascii").strip())
        if pid <= 0 or pid == os.getpid():
            return
        if not is_word_process(pid):
            return
        os.kill(pid, signal.SIGTERM)
    except (OSError, ValueError):
        pass


def is_word_process(pid: int) -> bool:
    if sys.platform != "win32":
        return False
    try:
        from ctypes import wintypes

        kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
        kernel32.OpenProcess.argtypes = [
            wintypes.DWORD,
            wintypes.BOOL,
            wintypes.DWORD,
        ]
        kernel32.OpenProcess.restype = wintypes.HANDLE
        kernel32.QueryFullProcessImageNameW.argtypes = [
            wintypes.HANDLE,
            wintypes.DWORD,
            wintypes.LPWSTR,
            ctypes.POINTER(wintypes.DWORD),
        ]
        kernel32.QueryFullProcessImageNameW.restype = wintypes.BOOL
        kernel32.CloseHandle.argtypes = [wintypes.HANDLE]
        kernel32.CloseHandle.restype = wintypes.BOOL

        handle = kernel32.OpenProcess(0x1000, False, pid)
        if not handle:
            return False
        try:
            buffer = ctypes.create_unicode_buffer(32768)
            size = wintypes.DWORD(len(buffer))
            if not kernel32.QueryFullProcessImageNameW(
                handle,
                0,
                buffer,
                ctypes.byref(size),
            ):
                return False
            executable = buffer.value
        finally:
            kernel32.CloseHandle(handle)
    except Exception:
        return False
    return Path(executable).name.lower() == "winword.exe"


def run_legacy_doc_parser(
    name: str,
    executable: str,
    path: Path,
) -> subprocess.CompletedProcess[bytes]:
    if name != "word-com":
        return subprocess.run(
            legacy_doc_command(name, executable, path),
            capture_output=True,
            check=False,
            timeout=120,
        )
    with tempfile.TemporaryDirectory(prefix="local-knowledge-word-com-") as tmp:
        pid_file = Path(tmp) / "word.pid"
        try:
            return subprocess.run(
                legacy_doc_command(name, executable, path, pid_file),
                capture_output=True,
                check=False,
                timeout=120,
            )
        finally:
            terminate_word_process(pid_file)


def parse_doc(path: Path) -> str:
    parsers = find_legacy_doc_parsers()
    if not parsers:
        raise LocalKnowledgeError(
            "DOC support requires a local legacy Word parser: "
            "Microsoft Word with pywin32 (Windows), textutil (macOS), "
            "antiword, or catdoc."
        )
    failures: List[str] = []
    for name, executable in parsers:
        try:
            result = run_legacy_doc_parser(name, executable, path)
        except subprocess.TimeoutExpired:
            failures.append(f"{name} timed out after 120 seconds")
            continue
        except OSError as exc:
            failures.append(f"{name} could not run: {exc}")
            continue
        if result.returncode != 0:
            detail = decode_text(result.stderr).strip()
            failure = f"{name} exited with status {result.returncode}"
            if detail:
                failure += f": {detail}"
            failures.append(failure)
            continue
        text = decode_text(result.stdout)
        if not text.strip():
            failures.append(f"{name} produced no text")
            continue
        return text
    raise LocalKnowledgeError(f"All available DOC parsers failed: {'; '.join(failures)}")


def html_to_text(raw: str) -> str:
    raw = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw)
    raw = re.sub(r"(?s)<[^>]+>", " ", raw)
    raw = html.unescape(raw)
    return re.sub(r"\n{3,}", "\n\n", raw)


def parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        raise LocalKnowledgeError("PDF support requires pypdf. Run doctor --install.") from exc
    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def parse_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:
        raise LocalKnowledgeError("DOCX support requires python-docx. Run doctor --install.") from exc
    doc = docx.Document(str(path))
    lines: List[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def parse_xlsx(path: Path) -> str:
    try:
        import openpyxl  # type: ignore
    except Exception as exc:
        raise LocalKnowledgeError("XLSX support requires openpyxl. Run doctor --install.") from exc
    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    lines: List[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                lines.append("\t".join(values))
    return "\n".join(lines)


def parse_csv_like(path: Path, delimiter: str) -> str:
    text = safe_read_text(path)
    rows: List[str] = []
    try:
        reader = csv.reader(text.splitlines(), delimiter=delimiter)
        for row in reader:
            rows.append("\t".join(row))
    except csv.Error:
        return text
    return "\n".join(rows)


def parse_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".doc":
        return parse_doc(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".xlsx":
        return parse_xlsx(path)
    if suffix == ".csv":
        return parse_csv_like(path, ",")
    if suffix == ".tsv":
        return parse_csv_like(path, "\t")
    text = safe_read_text(path)
    if suffix in {".html", ".htm", ".xml"}:
        return html_to_text(text)
    return text


def is_image_file(path: Path) -> bool:
    return path.suffix.lower() in IMAGE_EXTS


def optional_parser_error(path: Path) -> Optional[str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        try:
            import pypdf  # type: ignore  # noqa: F401
        except Exception:
            return "PDF support requires pypdf. Run doctor --install."
    if suffix == ".doc" and not find_legacy_doc_parser():
        return (
            "DOC support requires a local legacy Word parser: "
            "Microsoft Word with pywin32 (Windows), textutil (macOS), "
            "antiword, or catdoc."
        )
    if suffix == ".docx":
        try:
            import docx  # type: ignore  # noqa: F401
        except Exception:
            return "DOCX support requires python-docx. Run doctor --install."
    if suffix == ".xlsx":
        try:
            import openpyxl  # type: ignore  # noqa: F401
        except Exception:
            return "XLSX support requires openpyxl. Run doctor --install."
    return None


def supported_exts(mode: str) -> set[str]:
    exts = set(TEXT_EXTS) | set(OPTIONAL_EXTS)
    if mode == "multimodal":
        exts |= IMAGE_EXTS
    return exts


def should_skip_file(lowered_name: str, suffix: str) -> bool:
    if lowered_name in SKIP_FILE_NAMES:
        return True
    if suffix in SKIP_FILE_SUFFIXES:
        return True
    if lowered_name.endswith((".min.js", ".min.css", ".bundle.js", ".bundle.css")):
        return True
    if lowered_name.endswith((".generated.js", ".generated.ts", ".generated.tsx")):
        return True
    return False


def iter_files(root: Path, mode: str = DEFAULT_MODE) -> Iterable[Path]:
    allowed = supported_exts(mode)
    if root.is_file():
        suffix = root.suffix.lower()
        lowered = root.name.lower()
        if not should_skip_file(lowered, suffix) and suffix in allowed:
            yield root
        return
    for current, dirs, files in os.walk(root):
        dirs[:] = [d for d in dirs if d not in SKIP_DIRS and not d.startswith(".local-knowledge")]
        for name in sorted(files):
            path = Path(current) / name
            suffix = path.suffix.lower()
            lowered = name.lower()
            if should_skip_file(lowered, suffix):
                continue
            if suffix in allowed:
                yield path


def relative_doc_path(source: Path, file_path: Path) -> str:
    if source.is_file():
        return file_path.name
    return str(file_path.relative_to(source))


def normalize_space(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    return text.strip()


def line_offsets(text: str) -> List[int]:
    offsets = [0]
    for match in re.finditer("\n", text):
        offsets.append(match.end())
    return offsets


def line_for_offset(offsets: List[int], offset: int) -> int:
    lo, hi = 0, len(offsets)
    while lo + 1 < hi:
        mid = (lo + hi) // 2
        if offsets[mid] <= offset:
            lo = mid
        else:
            hi = mid
    return lo + 1


def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> List[Tuple[str, int, int]]:
    text = normalize_space(text)
    if not text:
        return []
    if chunk_overlap >= chunk_size:
        raise LocalKnowledgeError("--chunk-overlap must be smaller than --chunk-size.")
    chunks: List[Tuple[str, int, int]] = []
    offsets = line_offsets(text)
    start = 0
    length = len(text)
    while start < length:
        target_end = min(length, start + chunk_size)
        end = target_end
        if target_end < length:
            window = text[start:target_end]
            candidates = [window.rfind("\n\n"), window.rfind("\n"), window.rfind(". "), window.rfind("。")]
            best = max(candidates)
            if best >= max(200, chunk_size // 3):
                end = start + best + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append((chunk, line_for_offset(offsets, start), line_for_offset(offsets, end)))
        if end >= length:
            break
        start = max(0, end - chunk_overlap)
    return chunks


def scan_documents(
    source: Path,
    mode: str,
    chunk_size: int,
    chunk_overlap: int,
    max_file_mb: int,
    max_files: int,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]]]:
    files_meta: List[Dict[str, Any]] = []
    chunks: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []
    max_file_bytes = max_file_mb * 1024 * 1024
    accepted_files = 0
    for file_path in iter_files(source, mode):
        stat = None
        digest = None
        rel = None
        try:
            stat = file_path.stat()
            if accepted_files >= max_files:
                skipped.append({"path": str(file_path), "reason": f"max_files exceeded ({max_files})"})
                continue
            if stat.st_size > max_file_bytes:
                skipped.append({"path": str(file_path), "reason": f"file larger than --max-file-mb ({max_file_mb})"})
                continue
            if file_path.suffix.lower() in TEXT_EXTS and is_binary_sample(file_path):
                skipped.append({"path": str(file_path), "reason": "binary"})
                continue
            accepted_files += 1
            digest = sha256_file(file_path)
            rel = relative_doc_path(source, file_path)
            if is_image_file(file_path):
                if mode != "multimodal":
                    skipped.append({"path": str(file_path), "reason": "image requires --mode multimodal"})
                    continue
                file_chunks = [(f"[image] {rel}", 1, 1)]
            else:
                text = parse_file(file_path)
                file_chunks = chunk_text(text, chunk_size, chunk_overlap)
            file_meta = {
                "path": str(file_path.resolve()),
                "relative_path": rel,
                "sha256": digest,
                "mtime": stat.st_mtime,
                "size": stat.st_size,
                "modality": "image" if is_image_file(file_path) else "text",
                "chunks": len(file_chunks),
            }
            files_meta.append(file_meta)
            for idx, (chunk, line_start, line_end) in enumerate(file_chunks):
                chunk_id_seed = f"{rel}:{digest}:{idx}:{line_start}:{line_end}"
                chunk_id = hashlib.sha256(chunk_id_seed.encode("utf-8")).hexdigest()[:32]
                chunks.append({
                    "id": chunk_id,
                    "source": str(file_path.resolve()),
                    "relative_path": rel,
                    "chunk_index": idx,
                    "modality": "image" if is_image_file(file_path) else "text",
                    "line_start": line_start,
                    "line_end": line_end,
                    "text": chunk,
                })
        except Exception as exc:
            item: Dict[str, Any] = {
                "path": str(file_path.resolve()) if digest is not None else str(file_path),
                "reason": str(exc),
            }
            if stat is not None:
                item["mtime"] = stat.st_mtime
                item["size"] = stat.st_size
            if digest is not None:
                item["sha256"] = digest
            if rel is not None:
                item["relative_path"] = rel
            skipped.append(item)
    return files_meta, chunks, skipped


def scan_file_fingerprints(
    source: Path,
    mode: str,
    max_file_mb: int,
    max_files: int,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    files_meta: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []
    max_file_bytes = max_file_mb * 1024 * 1024
    accepted_files = 0
    for file_path in iter_files(source, mode):
        try:
            stat = file_path.stat()
            if accepted_files >= max_files:
                skipped.append({"path": str(file_path), "reason": f"max_files exceeded ({max_files})"})
                continue
            if stat.st_size > max_file_bytes:
                skipped.append({"path": str(file_path), "reason": f"file larger than --max-file-mb ({max_file_mb})"})
                continue
            if file_path.suffix.lower() in TEXT_EXTS and is_binary_sample(file_path):
                skipped.append({"path": str(file_path), "reason": "binary"})
                continue
            accepted_files += 1
            files_meta.append({
                "path": str(file_path.resolve()),
                "relative_path": relative_doc_path(source, file_path),
                "sha256": sha256_file(file_path),
                "mtime": stat.st_mtime,
                "size": stat.st_size,
                "modality": "image" if is_image_file(file_path) else "text",
            })
        except Exception as exc:
            skipped.append({"path": str(file_path), "reason": str(exc)})
    files_meta.sort(key=lambda item: item["relative_path"])
    skipped.sort(key=lambda item: item["path"])
    return files_meta, skipped
